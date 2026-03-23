from flask import Flask, render_template, request, send_file
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

app = Flask(__name__)

TEMPLATES = {
    'liquidation': {
        'title': 'Company Dissolution and Liquidation / تصفية وحل شركة',
        'arabic': [
            "قرر المالك الوحيد لشركة: {cn_ar} ، رخصة تجارية رقم: {tln} ، صادرة من دائرة الاقتصاد والسياحة بدبي، في مقر الشركة الرئيسي في دبي ما يلي:",
            "- حل وتصفية شركة ({cn_ar}) وفقا للقوانين والأنظمة المعمول بها في دولة الإمارات العربية المتحدة.",
            "- تعيين السادة : {acn_ar}، رخصة تجارية رقم: {actl} كمصفيين للشركة ليقوموا على الفور بحل وتصفية الشركة ومباشرة ومتابعة إجراءات التصفية . وللقيام بتقديم معاملة التصفية ومتابعة الإجراءات.",
            "ونظرا لعدم وجود اية أمور أخرى على جدول الاعمال فقد تقرر انهاء الاجتماع."
        ],
        'english': [
            "The single owner of {cn_en}, Commercial License No. ({tln}) issued by the Department of Economy and Tourism in Dubai, at the company’s headquarters in Dubai, includes the following:",
            "- Dissolution and liquidation of {cn_en}, in accordance with the laws and regulations in force in the United Arab Emirates.",
            "- Appointing {acn_en}, License No. {actl}, issued by Dubai Economy and Tourism, as liquidators of the company to immediately dissolve and liquidate the company and proceed to follow up the procedures for liquidating the company, in accordance with the applicable laws.",
            "and due to the absence of any other matters on the agenda, it was decided to end the meeting."
        ]
    }
}

@app.route('/')
def index():
    return render_template('index.html', templates=TEMPLATES)

@app.route('/generate', methods=['POST'])
def generate():
    template_key = request.form.get('template_key')
    selected = TEMPLATES[template_key]
    
    # Get all form data
    data = {
        'cn_ar': request.form.get('cn_ar'),
        'cn_en': request.form.get('cn_en'),
        'tln': request.form.get('tln'),
        'acn_ar': request.form.get('acn_ar'),
        'acn_en': request.form.get('acn_en'),
        'actl': request.form.get('actl')
    }
    
    doc = docx.Document()
    
    # Generate Arabic Section
    for p_text in selected['arabic']:
        p = doc.add_paragraph(p_text.format(**data))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    doc.add_paragraph("\n" + "_"*50 + "\n")
    
    # Generate English Section
    for p_text in selected['english']:
        doc.add_paragraph(p_text.format(**data))

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    
    return send_file(
        target_stream,
        as_attachment=True,
        download_name=f"Resolution_{data['cn_en']}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
app = Flask(__name__)

# The template dictionary
TEMPLATES = {
    'liquidation': {
        'title': 'Company Dissolution and Liquidation',
        'fields': ['cn', 'tln', 'acn', 'actl'],
        'arabic': [
            "قرر المالك الوحيد لشركة: {cn} ، رخصة تجارية رقم: {tln} ، صادرة من دائرة الاقتصاد والسياحة بدبي، في مقر الشركة الرئيسي في دبي ما يلي:",
            "- حل وتصفية شركة ({cn}) وفقا للقوانين والأنظمة المعمول بها في دولة الإمارات العربية المتحدة.",
            "- تعيين السادة : {acn}، رخصة تجارية رقم: {actl} كمصفيين للشركة ليقوموا على الفور بحل وتصفية الشركة ومباشرة ومتابعة إجراءات التصفية . وللقيام بتقديم معاملة التصفية ومتابعة الإجراءات.",
            "ونظرا لعدم وجود اية أمور أخرى على جدول الاعمال فقد تقرر انهاء الاجتماع."
        ],
        'english': [
            "The single owner of {cn}, Commercial License No. ({tln}) issued by the Department of Economy and Tourism in Dubai, at the company’s headquarters in Dubai, includes the following:",
            "- Dissolution and liquidation of {cn}, in accordance with the laws and regulations in force in the United Arab Emirates.",
            "- Appointing {acn}, License No. {actl}, issued by Dubai Economy and Tourism, as liquidators of the company to immediately dissolve and liquidate the company and proceed to follow up the procedures for liquidating the company, in accordance with the applicable laws.",
            "and due to the absence of any other matters on the agenda, it was decided to end the meeting."
        ]
    }
}

@app.route('/')
def index():
    return render_template('index.html', templates=TEMPLATES)

@app.route('/generate', methods=['POST'])
def generate():
    template_key = request.form.get('template_key')
    selected = TEMPLATES[template_key]
    
    # Collect data from form
    data = {field: request.form.get(field) for field in selected['fields']}
    
    # Create Document
    doc = docx.Document()
    
    # Arabic Section
    for p_text in selected['arabic']:
        p = doc.add_paragraph(p_text.format(**data))
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    doc.add_paragraph("\n" + "_"*50 + "\n")
    
    # English Section
    for p_text in selected['english']:
        doc.add_paragraph(p_text.format(**data))

    # Save to a byte stream so we don't need to store files on the server
    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    
    return send_file(
        target_stream,
        as_attachment=True,
        download_name=f"{template_key}_{data['cn']}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# --- 1. THE DICTIONARY TEMPLATE SYSTEM ---
# Add new transaction types here! Just use {cn}, {tln}, {acn}, and {actl} where needed.
TEMPLATES = {
    '1': {
        'title': 'Company Dissolution and Liquidation',
        'arabic': [
            "قرر المالك الوحيد لشركة: {cn} ، رخصة تجارية رقم: {tln} ، صادرة من دائرة الاقتصاد والسياحة بدبي، في مقر الشركة الرئيسي في دبي ما يلي:",
            "- حل وتصفية شركة ({cn}) وفقا للقوانين والأنظمة المعمول بها في دولة الإمارات العربية المتحدة.",
            "- تعيين السادة : {acn}، رخصة تجارية رقم: {actl} كمصفيين للشركة ليقوموا على الفور بحل وتصفية الشركة ومباشرة ومتابعة إجراءات التصفية . وللقيام بتقديم معاملة التصفية ومتابعة الإجراءات.",
            "ونظرا لعدم وجود اية أمور أخرى على جدول الاعمال فقد تقرر انهاء الاجتماع."
        ],
        'english': [
            "The single owner of {cn}, Commercial License No. ({tln}) issued by the Department of Economy and Tourism in Dubai, at the company’s headquarters in Dubai, includes the following:",
            "- Dissolution and liquidation of {cn}, in accordance with the laws and regulations in force in the United Arab Emirates.",
            "- Appointing {acn}, License No. {actl}, issued by Dubai Economy and Tourism, as liquidators of the company to immediately dissolve and liquidate the company and proceed to follow up the procedures for liquidating the company, in accordance with the applicable laws.",
            "and due to the absence of any other matters on the agenda, it was decided to end the meeting."
        ]
    },
    '2': {
         'title': 'Example: Change of Manager (Placeholder)',
         'arabic': ["قرر المالك لشركة {cn} تغيير المدير...", "رقم الرخصة: {tln}"],
         'english': ["The owner of {cn} decided to change the manager...", "License: {tln}"]
    }
}

def main():
    # --- 2. MENU SELECTION ---
    print("=== Document Generator ===")
    for key, template in TEMPLATES.items():
        print(f"{key}. {template['title']}")
    
    choice = input("\nSelect the transaction type (Enter the number): ")
    
    if choice not in TEMPLATES:
        print("Invalid choice. Exiting.")
        return
    
    selected_template = TEMPLATES[choice]

    # --- 3. DATA COLLECTION ---
    print(f"\n--- Entering data for: {selected_template['title']} ---")
    
    # Storing inputs in a dictionary so we can easily pass them to our text templates
    data = {
        'cn': input("Company Name (CN): "),
        'tln': input("Trade License Number (TLN): "),
        'acn': input("Auditing Company Name (ACN): "),
        'actl': input("Auditing Company Trade License (ACTL): ")
    }

    # --- 4. DOCUMENT GENERATION ---
    doc = docx.Document()

    # Generate Arabic (Right-Aligned)
    for paragraph_template in selected_template['arabic']:
        # .format(**data) replaces {cn} with the actual company name, etc.
        formatted_text = paragraph_template.format(**data) 
        p = doc.add_paragraph(formatted_text)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Separator
    doc.add_paragraph("\n" + "_"*50 + "\n")

    # Generate English (Left-Aligned)
    for paragraph_template in selected_template['english']:
        formatted_text = paragraph_template.format(**data)
        doc.add_paragraph(formatted_text)

    # --- 5. SAVE FILE (Routed to output folder for Docker) ---
    output_dir = "output"
    os.makedirs(output_dir, exist_ok=True) # Creates an 'output' folder if it doesn't exist
    
    # Create a safe filename without spaces
    safe_title = selected_template['title'].replace(" ", "_")
    safe_cn = data['cn'].replace(" ", "_")
    filename = os.path.join(output_dir, f"{safe_title}_{safe_cn}.docx")
    
    doc.save(filename)
    print(f"\nSuccess! Document saved as: {filename}")

if __name__ == "__main__":
    main()
